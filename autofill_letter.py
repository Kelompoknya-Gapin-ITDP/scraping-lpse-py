from docx import Document
from datetime import datetime

# Template function for creating the letter
def generate_letter(company_name, address, credit_facilities, nomor_surat):
    # Load the template document (assuming it's in the same directory)
    doc = Document()

    # Define today's date
    date_today = datetime.today().strftime('%d %B %Y')

    # Create the letter content with the provided details
    doc.add_paragraph(f"Bandung, {date_today}")
    doc.add_paragraph(f"Nomor       : {nomor_surat}")
    doc.add_paragraph("Sifat       : Penting")
    doc.add_paragraph("Lampiran    :")
    doc.add_paragraph("Perihal     : Penawaran Fasilitas Kredit")
    
    doc.add_paragraph(f"Kepada Yth:\nDirektur Utama\n{company_name}\n{address}\n")
    
    doc.add_paragraph(
        f"Dengan Hormat,\n\n"
        "Bersama ini kami seluruh Pimpinan dan Staf bank bjb menyampaikan salam hangat dan "
        "semoga sukses dalam menjalankan aktivitas sehari-hari. Dalam kesempatan ini, kami "
        f"bank bjb bermaksud menawarkan jalinan kerjasama dengan {company_name} "
        "dalam melayani aktivitas keuangan dan kebutuhan perbankan, yaitu kerjasama dalam "
        "hal pemberian fasilitas pembiayaan atau transaksi lainnya guna mendukung program "
        "serta meningkatkan kegiatan usaha Perusahaan.\n"
    )
    
    # Add credit facilities section based on user selection
    if 'Fasilitas Kredit Modal Kerja' in credit_facilities:
        doc.add_paragraph(
            "Fasilitas Kredit Modal Kerja\n"
            f"Calon Debitur   : {company_name}\n"
            "Plafond         : Maksimal Rp100.000.000.000,- (seratus miliar rupiah)\n"
            "Jenis Fasilitas Kredit  : Kredit Modal Kerja\n"
            "Tujuan Penggunaan       : Modal Kerja Usaha - Transaksional\n"
            "Sifat Kredit    : Revolving\n"
            "Suku Bunga Kredit: Minimal 11,00% p.a. efektif (negotiable)\n"
            "Provisi         : 0,5% dari plafon fasilitas kredit dan dibayarkan sekaligus pada saat "
            "penandatanganan Perjanjian Kredit. (negotiable)\n"
            "Jangka Waktu Kredit: Maksimal 12 (dua belas) bulan\n"
        )
    
    if 'Kredit Investasi' in credit_facilities:
        doc.add_paragraph(
            "Kredit Investasi\n"
            "Plafond Kredit  : Maksimal Rp100.000.000.000,- (seratus miliar rupiah)\n"
            "Jenis Kredit    : Kredit Investasi\n"
            "Sifat Kredit    : Non-revolving\n"
            "Tujuan Penggunaan: Ekspansi dan pengembangan Usaha\n"
            "Jangka Waktu Kredit: Maksimal sampai dengan 120 bulan\n"
            "Suku Bunga Kredit: Minimal 11,00% p.a. efektif (negotiable)\n"
            "Provisi         : 0,50% dari plafon fasilitas kredit (negotiable)\n"
        )
    
    if 'Garansi Bank' in credit_facilities:
        doc.add_paragraph(
            "Garansi Bank\n"
            "Jenis Fasilitas:\n"
            "Garansi Bank Penawaran\n"
            "Garansi Bank Pelaksanaan\n"
            "Garansi Bank Uang Muka\n"
            "Garansi Bank Pemeliharaan\n"
            "Garansi Bank Pembayaran\n"
            "Tujuan Penggunaan: Penerbitan garansi bank terkait pekerjaan dengan sumber dana APBN, APBD, BUMN, dan BUMD\n"
        )
    
    # Add the closing content
    doc.add_paragraph(
        "Persyaratan Kredit / Covenant\n"
        "a. Menyerahkan surat permohonan kredit yang ditandatangani oleh pejabat yang berwenang.\n"
        "b. Menyerahkan legalitas usaha dan perusahaan.\n"
        "c. Laporan keuangan\n"
        "d. Menyerahkan legalitas agunan.\n"
        "e. Dokumen dan/atau lainnya sesuai bidang usaha perusahaan dalam rangka proses analisa kredit\n\n"
        "Persyaratan Lainnya\n"
        "a. Proses kredit akan dilakukan analisa setelah seluruh data pengajuan kredit dipenuhi dan memenuhi seluruh ketentuan bank bjb.\n"
        "b. Tambahan persyaratan lain akan disesuaikan setelah proses analisa kredit dilaksanakan.\n\n"
        "Proses Pengajuan Kredit\n"
        "a. Proses pengajuan kredit akan diproses lebih lanjut setelah seluruh data/dokumen persyaratan pengajuan kredit dipenuhi dan memenuhi seluruh ketentuan bank bjb.\n"
        "b. Tambahan persyaratan lain akan disesuaikan dengan keputusan Komite Kredit.\n"
        "c. Penawaran ini hanya bersifat indikatif dan bukan merupakan komitmen pemberian fasilitas kredit serta tidak bersifat mengikat kedua belah pihak.\n\n"
        f"Besar harapan kami, kiranya {company_name} menggunakan produk/jasa perbankan kami sehingga dapat memberikan nilai tambah bagi kedua belah pihak dan kiranya data-data tersebut dapat kami terima pada kesempatan pertama dan dapat dikirimkan melalui email grupkomersial@outlook.com.\n\n"
        "Informasi lebih lanjut dapat menghubungi:\n"
        "Palupi Ratna Dewanti (081218796268)\n"
        "Anadia Nafila (081221809036)\n\n"
        "Demikian kami sampaikan, atas perkenan Bapak/Ibu kami ucapkan terima kasih.\n\n"
        "PT BANK PEMBANGUNAN DAERAH JAWA BARAT DAN BANTEN, Tbk.\n"
        "DIVISI KOMERSIAL\n\n"
        "Agus Somantri\n"
        "Pemimpin Divisi"
    )
    
    # Define the output file name
    output_filename = f"surat_penawaran_kredit_{nomor_surat}.docx"
    doc.save(output_filename)
    print(f"Document generated successfully as '{output_filename}'.")

# Get user inputs
company_name = input("Enter company name: ")
address = input("Enter company address: ")
credit_facilities = input("Enter credit facilities (comma separated): ").split(',')
nomor_surat = input("Enter the letter number (e.g., 611/KOM-KM1/2024): ")

# Generate the letter with user inputs
generate_letter(company_name.strip(), address.strip(), [cf.strip() for cf in credit_facilities], nomor_surat.strip())
